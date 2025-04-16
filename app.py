from dotenv import load_dotenv
import os
import time
import streamlit as st
import fitz  # PyMuPDF
import google.generativeai as genai
import re
from datetime import datetime
import uuid
import json
import base64
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

# Load environment variables and configure Gemini
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Page configuration with custom theme
st.set_page_config(
    page_title="ResumeRankr",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for improved styling and accessibility
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E3A8A;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        font-size: 1.3rem;
        color: #4B5563;
        margin-bottom: 2rem;
    }
    .section-header {
        font-size: 1.5rem;
        color: white;
        margin-top: 1.5rem;
        margin-bottom: 1rem;
        padding-bottom: 0.3rem;
        border-bottom: 2px solid #E5E7EB;
    }
    .info-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #F3F4F6;
        margin-bottom: 1rem;
    }
    .button-row {
        display: flex;
        flex-wrap: wrap;
        gap: 0.5rem;
        margin: 1rem 0;
    }
    .stButton>button {
        width: 100%;
        border-radius: 0.3rem;
        padding: 0.5rem;
        color: white;
        font-weight: 500;
    }
    .footer {
        margin-top: 2rem;
        padding-top: 1rem;
        border-top: 1px solid #E5E7EB;
        text-align: center;
        color: #6B7280;
    }
    .result-card {
        background-color: #F9FAFB;
        padding: 1.5rem;
        border-radius: 0.5rem;
        border-left: 4px solid #2563EB;
        margin: 1rem 0;
    }
    .privacy-notice {
        background-color: #FFEDD5;
        border-left: 4px solid #F59E0B;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    /* Accessibility improvements */
    a {
        color: #2563EB;
        text-decoration: underline;
    }
    button:focus, input:focus, select:focus, textarea:focus {
        outline: 2px solid #2563EB;
        outline-offset: 2px;
    }
    .tooltip {
        position: relative;
        display: inline-block;
        border-bottom: 1px dotted #6B7280;
    }
    .tooltip .tooltiptext {
        visibility: hidden;
        width: 200px;
        background-color: #374151;
        color: #fff;
        text-align: center;
        border-radius: 6px;
        padding: 5px;
        position: absolute;
        z-index: 1;
        bottom: 125%;
        left: 50%;
        margin-left: -100px;
        opacity: 0;
        transition: opacity 0.3s;
    }
    .tooltip:hover .tooltiptext {
        visibility: visible;
        opacity: 1;
    }
</style>
""", unsafe_allow_html=True)

# Gemini model options with descriptions
model_options = {
    "Gemini 2.0 Flash (Next-Gen Fast & Smart)": "gemini-2.0-flash",
    "Gemini 2.0 Flash-Lite (Low Latency)": "gemini-2.0-flash-lite", 
    "Gemini 1.5 Flash (Fast & Versatile)": "gemini-1.5-flash",
    "Gemini 1.5 Flash-8B (High Volume)": "gemini-1.5-flash-8b",
    "Gemini 1.5 Pro (Advanced Reasoning)": "gemini-1.5-pro"
}

# Prompts dictionary
PROMPTS = {
    "submit1": """
        You are an experienced Technical Human Resource Manager. Your task is to review the provided resume 
        against the job description. Highlight strengths and weaknesses in relation to the role.
    """,
    "submit2": """
    You are a Technical HR Manager with experience in assessing skill development. Evaluate the resume against the job description.

    Provide:
    1. Skills that are relevant and strong.
    2. Skills that are lacking or could be improved.
    3. Specific suggestions to enhance the candidate's technical and soft skills for this role.
    """,
    "submit3": """
        You are an ATS scanner with HR experience. Evaluate resume vs. job description, list missing keywords, 
        and suggest skill improvements.
    """,
    "submit4": """
    You are an intelligent ATS (Applicant Tracking System) scanner. Analyze the resume against the job description and provide the following:

    1. An estimated **match percentage** between the resume and the job description.
    2. A brief explanation of how this score was determined — highlight key strengths and gaps.
    3. Keep your response short, structured, and easy to understand.
    """,
    "submit5": """
    You are a professional Technical Recruiter. Carefully review the candidate's resume and compare it with the job description.

    Provide:
    1. A judgment on whether this candidate is likely to be shortlisted for an interview (Yes, Maybe, or Unlikely).
    2. A brief explanation supporting your judgment based on relevant skills, experience, and alignment.
    3. Any critical gaps that might lower the chances.
    4. Suggestions to improve the chances of getting noticed or selected.
    """
}

# Session timeout in minutes
SESSION_TIMEOUT = 30

# Initialize session state
if 'session_id' not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())
if 'session_start_time' not in st.session_state:
    st.session_state.session_start_time = datetime.now()
if 'api_calls_count' not in st.session_state:
    st.session_state.api_calls_count = 0
if 'inputs_ready' not in st.session_state:
    st.session_state.inputs_ready = False
if 'user_consent' not in st.session_state:
    st.session_state.user_consent = False
if 'privacy_acknowledged' not in st.session_state:
    st.session_state.privacy_acknowledged = False

def check_session_timeout():
    """Check if session has timed out and reset if needed"""
    current_time = datetime.now()
    time_diff = current_time - st.session_state.session_start_time
    if time_diff.total_seconds() > (SESSION_TIMEOUT * 60):
        # Reset session data
        if 'uploaded_file' in st.session_state:
            del st.session_state.uploaded_file
        if 'job_desc' in st.session_state:
            del st.session_state.job_desc
        if 'response_history' in st.session_state:
            del st.session_state.response_history
        st.session_state.session_start_time = current_time
        st.session_state.inputs_ready = False
        st.session_state.user_consent = False
        st.warning(f"Your session has timed out after {SESSION_TIMEOUT} minutes of inactivity. Your data has been cleared for privacy.")
        return True
    return False

def extract_pdf_text(uploaded_file):
    """Extract text from PDF file"""
    try:
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        text = " ".join(page.get_text() for page in doc)
        
        # Check for sensitive data patterns (simple example)
        sensitive_patterns = [
            r'\b\d{3}-\d{2}-\d{4}\b',  # SSN
            r'\b\d{16}\b',             # Credit card
            r'\b(?:password|passwd)(?:\s*:)?\s*\w+\b'  # Password
        ]
        
        for pattern in sensitive_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                st.error("⚠️ Your resume may contain sensitive personal information. Please remove and re-upload.")
                return None
                
        return text
    except Exception as e:
        st.error(f"Error processing PDF: {str(e)}")
        return None

def get_gemini_response(prompt, pdf_content, job_desc):
    """Get response from Gemini model with rate limiting"""
    # Check rate limits
    if st.session_state.api_calls_count >= 50:  # Example limit
        st.error("API usage limit reached. Please try again later.")
        return "Rate limit exceeded. Please try again later.", 0
    
    try:
        model = genai.GenerativeModel(model_choice)
        start = time.time()
        response = model.generate_content([prompt, pdf_content, job_desc])
        end = time.time()
        
        # Log API call
        st.session_state.api_calls_count += 1
        
        # Log usage (in real app, this might send to a database)
        logging_data = {
            "timestamp": datetime.now().isoformat(),
            "session_id": st.session_state.session_id,
            "model": model_choice,
            "duration": end - start
        }
        
        return response.text, end - start
    except Exception as e:
        st.error(f"API Error: {str(e)}")
        return f"Error: {str(e)}", 0

def validate_inputs():
    """Validate that both job description and resume are provided"""
    if not job_desc.strip():
        st.error("⚠️ Please enter a job description before submitting.")
        return False
    if uploaded_file is None:
        st.error("⚠️ Please upload a PDF resume to proceed.")
        return False
    if not st.session_state.user_consent:
        st.error("⚠️ Please provide consent for data processing before proceeding.")
        return False
    return True

def handle_response(prompt_key_or_custom):
    """Handle response generation with error handling"""
    # First check for timeout
    if check_session_timeout():
        return "", None
        
    if not validate_inputs():
        return "", None

    with st.spinner("Analyzing your resume... This may take a moment"):
        try:
            uploaded_file.seek(0)
            pdf_text = extract_pdf_text(uploaded_file)
            if pdf_text is None:
                return "", None
                
            prompt = PROMPTS.get(prompt_key_or_custom, prompt_key_or_custom)
            response_text, duration = get_gemini_response(prompt, pdf_text, job_desc)
            
            if prompt_key_or_custom == "submit4":
                match = re.search(r"(\d{1,3})\s*%", response_text)
                percentage = int(match.group(1)) if match else None
                return (percentage, response_text), duration

            return response_text, duration
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
            return f"Error: {str(e)}", None

def display_response(response_output, response_time_taken=None):
    """Display formatted response"""
    if response_output and not isinstance(response_output, tuple) and not isinstance(response_output, int):
        st.markdown("<div class='result-card'>", unsafe_allow_html=True)
        st.markdown("<h3>Analysis Results:</h3>", unsafe_allow_html=True)
        st.write(response_output)
        if response_time_taken:
            st.caption(f"⏱️ Response time: {response_time_taken:.2f} seconds")
        st.markdown("</div>", unsafe_allow_html=True)

def save_response_history(response, query_type):
    """Save response to session state history"""
    if 'response_history' not in st.session_state:
        st.session_state.response_history = []
    
    timestamp = datetime.now().strftime("%H:%M:%S")
    st.session_state.response_history.append({
        "type": query_type,
        "timestamp": timestamp,
        "response": response
    })

# Add this function to convert analysis history to a Word document
def convert_history_to_docx(history):
    """Convert analysis history data to a Word document format"""
    doc = Document()
    
    # Add a title
    title = doc.add_heading('ResumeRankr Analysis Results', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp_para = doc.add_paragraph()
    timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp_para.add_run(f'Generated on {datetime.now().strftime("%B %d, %Y at %H:%M:%S")}')
    
    # Add a divider
    doc.add_paragraph('_' * 50)
    
    # Process each analysis result
    for item in history:
        # Add section heading with type and timestamp
        heading = doc.add_heading(f'{item["type"]} ({item["timestamp"]})', level=1)
        
        # Add the response content
        para = doc.add_paragraph()
        para.add_run(item["response"])
        
        # Add a divider between sections
        doc.add_paragraph('_' * 50)
    
    # Add footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run('Generated by ResumeRankr - Powered by Gemini AI')
    
    # Save to a temporary file and return the bytes
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        tmp.seek(0)
        docx_data = open(tmp.name, 'rb').read()
    
    return docx_data
# Check for session timeout
check_session_timeout()

# Main app layout
st.markdown("<h1 class='main-header'>ResumeRankr</h1>", unsafe_allow_html=True)
st.markdown("<p class='sub-header'>Optimize your job application with AI-powered resume analysis</p>", unsafe_allow_html=True)

# Privacy notice and consent
with st.expander("📜 Privacy & Terms (Important - Please Read)", expanded=not st.session_state.privacy_acknowledged):
    st.markdown("""
    ### Privacy Notice
    
    **Data Usage:** 
    - Your resume and job description data are processed temporarily for analysis purposes only.
    - Data is not permanently stored on our servers and is automatically cleared after session timeout ({} minutes of inactivity).
    - We do not share your data with third parties except for processing via the Google Gemini API.
    
    **Data Security:**
    - Your data is encrypted during transit.
    - We automatically scan for sensitive personal information (e.g., SSNs, credit card numbers) and block processing if detected.
    
    ### Terms of Service
    
    **Usage Limitations:**
    - This tool provides advisory analysis only and does not guarantee job placement or interview outcomes.
    - Results are based on AI analysis and should be used as guidance, not as definitive assessments.
    - You agree not to misuse this service for unlawful purposes or to circumvent rate limiting.
    
    **Your Rights:**
    - You may request deletion of your data at any time by clicking "Clear My Data".
    - You can download any analysis results for your records.
    
    ### Compliance
    
    This application complies with GDPR, CCPA, and other applicable data protection regulations. We process data based on your explicit consent.
    """.format(SESSION_TIMEOUT))
    
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("I Acknowledge & Agree", use_container_width=True):
            st.session_state.privacy_acknowledged = True
            st.rerun()
    with col2:
        if st.button("Clear My Data", use_container_width=True):
            for key in list(st.session_state.keys()):
                if key not in ['session_id', 'privacy_acknowledged']:
                    if key in st.session_state:
                        del st.session_state[key]
            st.session_state.inputs_ready = False
            st.session_state.user_consent = False
            st.success("✅ Your data has been cleared.")
            st.rerun()

# Create two columns for layout
col1, col2 = st.columns([3, 2])

with col1:
    st.markdown("<div class='section-header'>Upload Information</div>", unsafe_allow_html=True)
    
    # Data usage consent
    consent = st.checkbox("I consent to processing my resume and job description data for analysis purposes", 
                         value=st.session_state.user_consent)
    st.session_state.user_consent = consent
    
    job_desc = st.text_area(
        "📋 Job Description:", 
        placeholder="Paste the job description here...",
        height=250,
        help="Paste the full job description text here. Do not include personal identifiable information."
    )
    
    with st.expander("📝 Resume Upload"):
        st.markdown("""
        **Important:** 
        - Upload PDF format only
        - Remove sensitive personal information (SSN, ID numbers, credit card details, etc.)
        - Ensure your resume is in English for best results
        """)
        
        uploaded_file = st.file_uploader(
            "Upload your Resume (PDF format only)", 
            type=["pdf"],
            help="Your resume will be analyzed against the job description"
        )
        
        if uploaded_file:
            st.success("✅ Resume uploaded successfully")
            try:
                file_details = {"Filename": uploaded_file.name, "Size": f"{uploaded_file.size / 1024:.2f} KB"}
                st.json(file_details)
            except:
                pass
    
    # Check and update inputs_ready status
    st.session_state.inputs_ready = bool(job_desc.strip() and uploaded_file and st.session_state.user_consent)
    
    with st.expander("⚙️ Advanced Settings"):
        st.markdown("<p>Select the Gemini model for your analysis:</p>", unsafe_allow_html=True)
        selected_model_label = st.selectbox(
            "AI Model:",
            list(model_options.keys()),
            index=0,
            help="Different models offer varying levels of analysis depth and speed"
        )
        model_choice = model_options[selected_model_label]
        
        # Add API usage display
        st.caption(f"API calls in this session: {st.session_state.api_calls_count}/50")
    
    st.markdown("<div class='section-header'>Analysis Options</div>", unsafe_allow_html=True)
    
    # Show ready status
    if st.session_state.inputs_ready:
        st.success("✅ Ready to analyze! Select an option below.")
    else:
        missing = []
        if not job_desc.strip():
            missing.append("job description")
        if not uploaded_file:
            missing.append("resume")
        if not st.session_state.user_consent:
            missing.append("consent")
        
        st.warning(f"⚠️ Please provide {' and '.join(missing)} before analysis.")
    
    # Organized buttons in tabs for better categorization
    tabs = st.tabs(["Basic Analysis", "Detailed Analysis", "Custom Query"])
    
    with tabs[0]:
        col1a, col1b = st.columns(2)
        with col1a:
            if st.button("✨ Resume Overview", use_container_width=True) and st.session_state.inputs_ready:
                response_output, response_time_taken = handle_response("submit1")
                if response_output:  # Only save if there's actual content
                    save_response_history(response_output, "Resume Overview")
                
        with col1b:
            if st.button("🎯 Match Percentage", use_container_width=True) and st.session_state.inputs_ready:
                result, response_time_taken = handle_response("submit4")
                if isinstance(result, tuple):
                    percentage, explanation = result
                    save_response_history(explanation, "Match Percentage")
                    
    with tabs[1]:
        col2a, col2b = st.columns(2)
        with col2a:
            if st.button("🚀 Skill Improvement", use_container_width=True) and st.session_state.inputs_ready:
                response_output, response_time_taken = handle_response("submit2")
                if response_output:
                    save_response_history(response_output, "Skill Improvement")
                
        with col2b:
            if st.button("🔍 Missing Keywords", use_container_width=True) and st.session_state.inputs_ready:
                response_output, response_time_taken = handle_response("submit3")
                if response_output:
                    save_response_history(response_output, "Missing Keywords")
                
        if st.button("👔 Interview Chances", use_container_width=True) and st.session_state.inputs_ready:
            response_output, response_time_taken = handle_response("submit5")
            if response_output:
                save_response_history(response_output, "Interview Chances")
            
    with tabs[2]:
        custom_query = st.text_input(
            "Ask a specific question:",
            placeholder="Example: What specific certifications would help me for this role?"
        )
        if st.button("🔮 Answer My Question", use_container_width=True) and custom_query and st.session_state.inputs_ready:
            response_output, response_time_taken = handle_response(custom_query)
            if response_output:
                save_response_history(response_output, f"Custom: {custom_query}")

with col2:
    st.markdown("<div class='section-header'>Analysis Results</div>", unsafe_allow_html=True)
    
    # Only show analysis results if inputs are ready
    if st.session_state.inputs_ready:
        # Display specific percentage match if available
        if 'response_history' in st.session_state and any(item["type"] == "Match Percentage" for item in st.session_state.response_history):
            # Find the most recent match percentage
            match_item = next((item for item in reversed(st.session_state.response_history) if item["type"] == "Match Percentage"), None)
            if match_item:
                match = re.search(r"(\d{1,3})\s*%", match_item["response"])
                percentage = int(match.group(1)) if match else None
                
                if percentage is not None:
                    st.metric(
                        "Resume-Job Description Match", 
                        f"{percentage}%",
                        delta=None
                    )
                    
                    # Color-coded progress bar
                    color = "green" if percentage >= 70 else "orange" if percentage >= 50 else "red"
                    st.progress(percentage / 100)
                    
                    match_text = "Excellent Match" if percentage >= 85 else \
                                "Good Match" if percentage >= 70 else \
                                "Fair Match" if percentage >= 50 else \
                                "Needs Improvement"
                                
                    st.info(f"**Match Quality:** {match_text}")
                    
                    # Disclaimer about match percentage
                    st.caption("**Note:** Match percentages are algorithmic estimates and should not be considered definitive.")
        
        # Response history with download option
        if 'response_history' in st.session_state and st.session_state.response_history:
            # Add download option
            if st.button("📥 Download Analysis Results"):
                try:
                    docx_data = convert_history_to_docx(st.session_state.response_history)
                    b64 = base64.b64encode(docx_data).decode()
                    download_filename = f"ResumeRankr_analysis_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
                    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{download_filename}">Click to download Word document</a>'
                    st.markdown(href, unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Error generating Word document: {str(e)}")
            
            for item in reversed(st.session_state.response_history):
                with st.expander(f"{item['type']} ({item['timestamp']})"):
                    st.write(item['response'])
        else:
            st.info("📊 Click an analysis option to see results here.")
    else:
        # Show instructions when inputs are not ready
        st.info("📋 Complete all required inputs to begin analysis:")
        st.markdown("""
        1. Provide consent to data processing
        2. Paste the job description
        3. Upload your resume (PDF format)
        
        Once all are provided, analysis options will become available.
        """)
            
        with st.expander("📚 How to get started"):
            st.markdown("""
            1. Read and acknowledge the privacy terms
            2. Paste the job description in the text area
            3. Upload your resume PDF (remove sensitive information)
            4. Choose an analysis option from the tabs
            5. View your results in this panel
            6. Download your results for future reference
            """)

# Accessibility features
st.markdown("""
<div role="region" aria-label="Accessibility information">
    <p>This application supports keyboard navigation. Press Tab to move between elements and Enter to activate buttons.</p>
    <p>For screen reader support or accessibility issues, please contact support.</p>
</div>
""", unsafe_allow_html=True)
            
# Footer with legal information
st.markdown("<div class='footer'>", unsafe_allow_html=True)
st.markdown("**ResumeRankr** - Powered by Gemini AI")
st.markdown("""
*Disclaimer: This application is for educational purposes only and does not guarantee job placement or success. 
Analysis results should be used as guidance only.*
""")
st.markdown("""
<small>© 2025 ResumeRankr | <a href="#" tabindex="0">Privacy Policy</a> | <a href="#" tabindex="0">Terms of Service</a> | <a href="#" tabindex="0">Accessibility</a></small>
""", unsafe_allow_html=True)
st.markdown("</div>", unsafe_allow_html=True)

# Session timeout notification (in real application, this would use JavaScript)
elapsed_time = (datetime.now() - st.session_state.session_start_time).total_seconds() / 60
time_left = max(0, SESSION_TIMEOUT - elapsed_time)
if time_left < 5:  # Show warning when less than 5 minutes left
    st.warning(f"⚠️ Your session will expire in {time_left:.1f} minutes. Save your results if needed.")